#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_cv.py — Reusable tailored-CV renderer for Roei Klein's job-hunt pipeline.

Used by the job-hunt-agent (autonomous) and the assisted-apply flow.
Takes a JSON spec of the TAILORED content and renders a polished .docx.

Usage:
  python3 render_cv.py --data spec.json --out /path/CV_Company_Role.docx
  echo '<json>' | python3 render_cv.py --out /path/out.docx        # JSON via stdin

JSON schema:
{
  "summary": "tailored professional summary paragraph",
  "experience": [
     {"title": "Role | Company", "dates": "MM/YYYY - Present", "bullets": ["XYZ bullet", ...]},
     ...
  ],
  "education": ["line 1", "line 2"],          # optional, sensible defaults if omitted
  "skills": [["Label","comma, separated, items"], ...]
}
Identity/contact are fixed (Roei) but can be overridden via "name","headline","contact".
Experience dates are auto-normalized to ATS-safe "MM/YYYY - MM/YYYY" / "MM/YYYY - Present".
"""
import sys, os, json, argparse, subprocess, re

try:
    from docx import Document
except ImportError:
    subprocess.run([sys.executable,"-m","pip","install","python-docx","--break-system-packages","-q"], check=False)
    from docx import Document

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x33,0x55); GREY=RGBColor(0x55,0x55,0x55)

# --- ATS date normalization ------------------------------------------------
# ATS engines compute Years-of-Experience from the Experience dates. A Jobscan
# study found ~37% of resumes were down-ranked because dates could not be parsed.
# Enforce one consistent, parseable format: "MM/YYYY - MM/YYYY" / "MM/YYYY - Present".
# Year-only inputs default to month 01 (start) / 12 (end) to keep the format uniform.
_MONTHS={'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12}
def _norm_token(t, is_end):
    t=(t or "").strip()
    if not t: return t
    if re.match(r'(?i)^(present|current|now|ongoing|today)$', t): return "Present"
    m=re.match(r'^(\d{1,2})[/.\-](\d{4})$', t)
    if m: return f"{int(m.group(1)):02d}/{m.group(2)}"
    m=re.match(r'(?i)^([a-z]{3})[a-z]*\.?\s+(\d{4})$', t)
    if m and m.group(1).lower() in _MONTHS: return f"{_MONTHS[m.group(1).lower()]:02d}/{m.group(2)}"
    m=re.match(r'^(\d{4})$', t)
    if m: return f"{'12' if is_end else '01'}/{m.group(1)}"
    return t
def normalize_dates(s):
    """Return a range normalized to 'MM/YYYY - MM/YYYY' / 'MM/YYYY - Present'."""
    if not s: return s
    parts=re.split(r'\s*[–—\-]\s*', s.strip(), maxsplit=1)
    if len(parts)==2:
        return f"{_norm_token(parts[0], False)} - {_norm_token(parts[1], True)}"
    return _norm_token(s, False)
# ---------------------------------------------------------------------------

DEFAULT_EDU=[
 "M.Sc. Data Science & Artificial Intelligence | Bar-Ilan University, 2025-2027 - Coursework: Machine Learning, Advanced Python, SQL, Statistical Modeling, AI Applications.",
 "B.A. Economics & Business Administration | Bar-Ilan University, 2022-2025 - GPA: 88, Graduated with Distinction.",
]

def underline_para(p, sz='6', space='2'):
    """Attach the rule directly to the heading paragraph (no separate empty line),
    so the title and its rule read as ONE element instead of a detached bar."""
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr')
    bo=OxmlElement('w:bottom')
    for k,v in (('w:val','single'),('w:sz',sz),('w:space',space),('w:color','1F3355')): bo.set(qn(k),v)
    pbdr.append(bo); pPr.append(pbdr)

def build(spec, out):
    name=spec.get("name","ROEI KLEIN")
    headline=spec.get("headline","Data Analyst  |  Product & Growth Oriented  |  Data Science Background")
    contact=spec.get("contact","Hod HaSharon, Israel  •  tklroei1@gmail.com  •  054-3329092  •  linkedin.com/in/roei-klein")
    # compact=True → squeeze the same content onto one page (tighter leading/margins,
    # slightly smaller body font). Content is NOT cut; only whitespace is.
    compact=bool(spec.get("compact"))
    BODY = Pt(spec.get("body_pt", 9.5 if compact else 10.5))  # body_pt lets you keep full-size text with tight leading
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=BODY
    st.paragraph_format.space_after=Pt(1 if compact else 6)
    st.paragraph_format.space_before=Pt(0)
    st.paragraph_format.line_spacing=1.0 if compact else 1.08
    for s in doc.sections:
        m=0.4 if compact else 0.6
        s.top_margin=Inches(m); s.bottom_margin=Inches(m); s.left_margin=Inches(0.55 if compact else 0.7); s.right_margin=Inches(0.55 if compact else 0.7)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(0)
    r=p.add_run(name); r.bold=True; r.font.size=Pt(17 if compact else 20); r.font.color.rgb=NAVY
    for txt,sz in ((headline,9.5 if compact else 10.5),(contact,9 if compact else 9.5)):
        q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after=Pt(0)
        rr=q.add_run(txt); rr.font.size=Pt(sz); rr.font.color.rgb=GREY
    def heading(t):
        h=doc.add_paragraph()
        h.paragraph_format.space_before=Pt(7 if compact else 9); h.paragraph_format.space_after=Pt(3 if compact else 4)
        h.paragraph_format.line_spacing=1.0
        rr=h.add_run(t.upper()); rr.bold=True; rr.font.size=Pt(10.5 if compact else 11.5)
        rr.font.color.rgb=NAVY; rr.font.name='Calibri'
        rPr=rr._element.get_or_add_rPr(); sp_=OxmlElement('w:spacing'); sp_.set(qn('w:val'),'24'); rPr.append(sp_)
        underline_para(h)

    def bullet(t):
        b=doc.add_paragraph(style='List Bullet'); b.paragraph_format.space_after=Pt(0 if compact else 2)
        b.paragraph_format.line_spacing=1.0 if compact else 1.08
        b.add_run(t).font.size=BODY

    def entries(items):
        for role in items:
            rp=doc.add_paragraph(); rp.paragraph_format.space_after=Pt(0)
            rp.paragraph_format.space_before=Pt(3 if compact else 4)
            rp.paragraph_format.line_spacing=1.0
            t=rp.add_run(role["title"]); t.bold=True; t.font.size=BODY
            if role.get("dates"):
                d=rp.add_run("\t"+normalize_dates(role["dates"])); d.italic=True
                d.font.size=Pt(9.5); d.font.color.rgb=GREY
                rp.paragraph_format.tab_stops.add_tab_stop(Inches(7.1 if compact else 7.0), WD_TAB_ALIGNMENT.RIGHT)
            for bl in role.get("bullets",[]): bullet(bl)

    def skills_block():
        for label, items in spec.get("skills",[]):
            b=doc.add_paragraph(); b.paragraph_format.space_after=Pt(0 if compact else 1)
            b.paragraph_format.line_spacing=1.0 if compact else 1.08
            r1=b.add_run(label+":  "); r1.bold=True; r1.font.size=BODY
            r2=b.add_run(items); r2.font.size=BODY

    def section(key):
        if key=="summary":
            heading("Professional Summary")
            sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(1 if compact else 6)
            sp.paragraph_format.line_spacing=1.0 if compact else 1.08
            sp.add_run(spec["summary"]).font.size=BODY
        elif key=="education":
            heading("Education")
            for line in spec.get("education",DEFAULT_EDU): bullet(line)
        elif key=="experience":
            heading("Experience"); entries(spec.get("experience",[]))
        elif key=="projects":
            if spec.get("projects"):
                heading("Projects"); entries(spec["projects"])
        elif key=="skills":
            heading("Skills"); skills_block()

    # Default order: Education before Experience (student positioning); Projects get their own section.
    for key in spec.get("order", ["summary","education","experience","projects","skills"]):
        section(key)
    os.makedirs(os.path.dirname(os.path.abspath(out)), exist_ok=True)
    doc.save(out); return out

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument("--data", help="path to JSON spec (or use stdin)")
    ap.add_argument("--out", required=True)
    a=ap.parse_args()
    raw=open(a.data,encoding="utf-8").read() if a.data else sys.stdin.read()
    spec=json.loads(raw)
    print("saved:", build(spec, a.out))

if __name__=="__main__":
    main()
